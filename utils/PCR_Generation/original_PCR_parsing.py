import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import docx
import json
import os

SPECIAL_HEADERS = [
        "Obligation",
        "License",
        "Release",
        "Status",
        "Type",
        "Comment"
    ]

def parse_docx_to_hierarchical_json(docx_path, output_path=None, verbose=False):
    """
    将Word文档解析为层次化的JSON结构
    
    参数:
    docx_path - Word文档路径
    output_path - 输出JSON文件路径（可选）
    verbose - 是否打印详细信息
    
    返回:
    层次化的JSON结构，包含文档元数据、章节结构和表格
    """
    if verbose:
        print(f"开始解析文档: {docx_path}")
    
    try:
        # 打开文档
        doc = docx.Document(docx_path)
        
        # 创建结果结构
        result = {
            "metadata": {
                "filename": os.path.basename(docx_path),
                "paragraphs_count": len(doc.paragraphs),
                "tables_count": len(doc.tables) if hasattr(doc, "tables") else 0
            },
            "structure": {
                "sections": []
            },
            "tables": []
        }
        
        # 处理章节结构
        current_sections = [None] * 10  # 支持10级标题
        
        # 遍历段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            # 检查是否为标题
            heading_level = 0
            if para.style and para.style.name.startswith('Heading'):
                try:
                    heading_level = int(para.style.name.replace('Heading', ''))
                except ValueError:
                    heading_level = 0
            
            # 如果段落是加粗的，可能是自定义标题
            is_bold = any(run.bold for run in para.runs if run.text.strip())
            if heading_level == 0 and is_bold:
                heading_level = 1  # 将加粗文本视为一级标题
            
            # 处理标题段落
            if heading_level > 0:
                # 创建新章节
                new_section = {
                    "title": text,
                    "level": heading_level,
                    "content": [],
                    "subsections": []
                }
                
                # 放置在正确的位置
                if heading_level == 1:
                    # 顶级章节
                    result["structure"]["sections"].append(new_section)
                    current_sections[0] = new_section
                    # 清除更深层次
                    for i in range(1, 10):
                        current_sections[i] = None
                else:
                    # 查找父章节
                    parent_level = heading_level - 1
                    parent_section = current_sections[parent_level]
                    
                    if parent_section:
                        parent_section["subsections"].append(new_section)
                        current_sections[heading_level] = new_section
                        # 清除更深层次
                        for i in range(heading_level + 1, 10):
                            current_sections[i] = None
                    else:
                        # 找不到父章节，作为顶级章节处理
                        result["structure"]["sections"].append(new_section)
                        current_sections[0] = new_section
                        current_sections[heading_level] = new_section
            else:
                # 普通段落文本
                # 查找当前章节
                current_section = None
                for i in range(9, -1, -1):
                    if current_sections[i]:
                        current_section = current_sections[i]
                        break
                
                if current_section:
                    # 添加到当前章节
                    current_section["content"].append(text)
                else:
                    # 未分类文本
                    if "uncategorized" not in result["structure"]:
                        result["structure"]["uncategorized"] = []
                    result["structure"]["uncategorized"].append(text)
        
        # 处理表格
        for i, table in enumerate(doc.tables):
            try:
                # 提取表头和数据
                headers = []
                data_rows = []
                
                if len(table.rows) > 0:
                    # 获取表头
                    header_row = table.rows[0]
                    headers = [cell.text.strip() for cell in header_row.cells]
                    
                    # 生成默认表头（如果为空）
                    for j in range(len(headers)):
                        if not headers[j]:
                            headers[j] = f"Column_{j+1}"
                    
                    is_special_table = (headers == SPECIAL_HEADERS)

                    if is_special_table:
                        row_idx = 1
                        while row_idx < len(table.rows) - 1 :
                            row = table.rows[row_idx]
                            next_row = table.rows[row_idx + 1]

                            row_data = {}
                            for col_idx, cell in enumerate(row.cells):
                                if col_idx < len(headers):
                                    row_data[headers[col_idx]] = cell.text.strip()

                            description_text = ' '.join([cell.text.strip() for cell in next_row.cells])
                            row_data['description'] = description_text

                            if any(v for k,v in row_data.items() if k != 'description'):
                                data_rows.append(row_data)

                            row_idx += 2

                    else:
                        # 处理数据行
                        for row_idx in range(1, len(table.rows)):
                            row = table.rows[row_idx]
                            row_data = {}
                            
                            for col_idx, cell in enumerate(row.cells):
                                if col_idx < len(headers):
                                    row_data[headers[col_idx]] = cell.text.strip()
                            
                            # 只添加非空行
                            if any(row_data.values()):
                                data_rows.append(row_data)
                
                # 创建表格结构
                table_structure = {
                    "id": f"table_{i+1}",
                    "title": f"Table {i+1}",
                    "headers": headers,
                    "data": data_rows
                }
                
                result["tables"].append(table_structure)
                
            except Exception as e:
                # 处理表格解析错误
                if verbose:
                    print(f"处理表格 #{i+1} 时出错: {str(e)}")
                result["tables"].append({
                    "id": f"table_{i+1}",
                    "error": str(e)
                })
        
        # 保存结果
        if output_path:
            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(result, f, ensure_ascii=False, indent=2)
            if verbose:
                print(f"结果已保存到: {output_path}")
        
        return result
        
    except Exception as e:
        error_msg = f"解析文档时出错: {str(e)}"
        if verbose:
            print(error_msg)
        
        if output_path:
            error_result = {"error": error_msg}
            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(error_result, f, ensure_ascii=False, indent=2)
        
        return {"error": error_msg}

# 示例使用
if __name__ == "__main__":
    doc_path = r"uploads\test\ProjectClearingReport-Wireless Room Sensor-2.0-2025-08-28_03_14_37.docx"

    json_output =parse_docx_to_hierarchical_json(doc_path, 'uploads/test/test.json')
    print(json_output)